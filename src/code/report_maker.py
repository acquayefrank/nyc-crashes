import pandas as pd
from datetime import datetime, timedelta
import os
import json
import sqlite3 as sql
from docx import Document
from docx.shared import Inches
from matplotlib import pyplot as plt
import numpy as np
from io import BytesIO
from src.code.predict_maker import build_predictions
import logging

from src.code.constants import (
    CURRENT_FILENAMES_FILE, METABASE_DATA_PATH, UTILS_PATH,
    CRASHES_AFTER_ETL, SELECT_ALL_QUERY, DATETIME, DAY_FORMAT,
    REPORT_FOR_CITIZENS_FILE, REPORT_FOR_MAYOR_FILE, RESULT_FOLDER
)

def make_report() -> None:
    """
    Makes reports for citizens and mayor.
    Reports are saved to REPORT_FOR_CITIZENS_FILE and REPORT_FOR_MAYOR_FILE respectively.

    :return:
    """

    with open(os.path.join(UTILS_PATH, CURRENT_FILENAMES_FILE), "r") as f:
        FILENAMES = json.load(f)

    logging.info("Reports are being created.")

    filename = FILENAMES[CRASHES_AFTER_ETL]
    conn = sql.connect(f"{METABASE_DATA_PATH}{filename}.db")

    query = f"{SELECT_ALL_QUERY} '{filename}'"

    df = pd.read_sql_query(query, conn)
    df[DATETIME] = df[DATETIME].apply(lambda x: x.split("T")[0])

    yesterday = datetime.today() - timedelta(days=5) # shift for 5 days
    last_week = set([(yesterday - timedelta(days=x)).strftime(DAY_FORMAT) for x in range(7)])
    last_month = set([(yesterday - timedelta(days=x)).strftime(DAY_FORMAT) for x in range(30)])
    yesterday = yesterday.strftime(DAY_FORMAT)

    yesterday_data = df[df[DATETIME] == yesterday]
    last_week_data = df[df[DATETIME].isin(last_week)]
    last_month_data = df[df[DATETIME].isin(last_month)]

    today_for_report = datetime.today().strftime("%d %B %Y")
    yesterday_for_report = (datetime.today() - timedelta(days=1)).strftime("%d %B %Y")
    last_week_for_report = (datetime.today() - timedelta(days=7)).strftime("%d %B %Y")
    last_month_for_report = (datetime.today() - timedelta(days=30)).strftime("%d %B %Y")

    names = [f"Yesterday ({yesterday_for_report})",
             f"Last Week ({last_week_for_report} - {yesterday_for_report})",
             f"Last Month ({last_month_for_report} - {yesterday_for_report})"]

    datasets = [yesterday_data, last_week_data, last_month_data]

    document = Document()

    document.add_heading("Crash Report", 0)
    document.add_heading(f"Report is generated on {today_for_report}", 1)

    for name, data in zip(names, datasets):
        document.add_paragraph()

        injured = int(data["persons_injured"].sum())
        killed = int(data["persons_killed"].sum())
        street_info = data["street_name"].value_counts()[:10].to_dict()

        document.add_heading(f"{name}:", 2)
        document.add_paragraph(f"Number of accidents: {len(data)}", style="List Bullet")
        document.add_paragraph(f"Number of people injured: {injured}", style="List Bullet")
        document.add_paragraph(f"Number of people killed: {killed}", style="List Bullet")

        p = document.add_heading("Streets with the highest number of accidents", 3)
        p.alignment = 1

        memfile = BytesIO()
        plt.figure(figsize=(6, 4))
        plt.barh(np.arange(9, -1, -1), street_info.values())
        plt.yticks(np.arange(9, -1, -1), street_info.keys())
        plt.ylabel("Street Name")
        plt.xlabel("Number of Accidents")
        plt.grid(axis="x")
        plt.tight_layout()
        plt.savefig(memfile)

        document.add_picture(memfile, width=Inches(6))

    document.save(os.path.join(RESULT_FOLDER, REPORT_FOR_CITIZENS_FILE))
    logging.info("Report for citizens successfully created.")

    streets_predicted = build_predictions()
    document.add_heading("List of streets where accidents are more likely to occur tomorrow:", 2)

    for street in streets_predicted:
        document.add_paragraph(street, style="List Number")

    document.save(os.path.join(RESULT_FOLDER, REPORT_FOR_MAYOR_FILE))
    logging.info("Report for mayor successfully created.")